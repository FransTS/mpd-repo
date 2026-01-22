"""
Output Generator for Transcription Results
Generates JSON and DOCX formats

Author: Frans Vermaak, CTGO @ LarcAI
"""

import json
import logging
from datetime import datetime
from pathlib import Path
from typing import Optional

logger = logging.getLogger(__name__)


def format_timestamp(seconds: float) -> str:
    """Convert seconds to HH:MM:SS format"""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    else:
        return f"{minutes:02d}:{secs:02d}"


class OutputGenerator:
    """
    Generates output files from transcription results.
    
    Formats:
    - JSON: Structured data with all metadata
    - DOCX: Professional document with speaker colours
    """
    
    # Speaker colours for DOCX (RGB tuples)
    SPEAKER_COLOURS = [
        (0, 51, 102),     # Dark blue
        (102, 51, 0),     # Brown
        (0, 102, 51),     # Dark green
        (102, 0, 102),    # Purple
        (153, 76, 0),     # Orange-brown
        (0, 102, 102),    # Teal
    ]
    
    def __init__(self):
        pass
    
    def generate_json(self, result: dict, output_dir: Path) -> Path:
        """
        Generate JSON output file.
        
        Args:
            result: Transcription result
            output_dir: Output directory
        
        Returns:
            Path to generated file
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_path = output_dir / "transcript.json"
        
        # Build output structure
        output = {
            "metadata": {
                "generated_at": datetime.now().isoformat(),
                "duration_seconds": result.get("duration", 0),
                "duration_formatted": format_timestamp(result.get("duration", 0)),
                "language": result.get("language", "auto"),
                "model": result.get("model", "unknown"),
                "segment_count": len(result.get("segments", [])),
                "speaker_count": len(result.get("speakers", [])),
                "vocabulary_corrections": result.get("vocabulary_corrections", 0)
            },
            "speakers": result.get("speakers", []),
            "segments": []
        }
        
        # Process segments
        for seg in result.get("segments", []):
            segment = {
                "start": seg.get("start", 0),
                "end": seg.get("end", 0),
                "start_formatted": format_timestamp(seg.get("start", 0)),
                "end_formatted": format_timestamp(seg.get("end", 0)),
                "text": seg.get("text", ""),
                "speaker": seg.get("speaker_name", seg.get("speaker", "UNKNOWN")),
            }
            
            # Add original text if corrected
            if "text_original" in seg:
                segment["text_original"] = seg["text_original"]
            
            # Add word timestamps if available
            if "words" in seg:
                segment["words"] = seg["words"]
            
            output["segments"].append(segment)
        
        # Write JSON
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(output, f, indent=2, ensure_ascii=False)
        
        logger.info(f"Generated JSON: {output_path}")
        return output_path
    
    def generate_docx(
        self,
        result: dict,
        output_dir: Path,
        include_timestamps: bool = True,
        include_word_timestamps: bool = False
    ) -> Path:
        """
        Generate DOCX output file.
        
        Args:
            result: Transcription result
            output_dir: Output directory
            include_timestamps: Include segment timestamps
            include_word_timestamps: Include word-level timestamps
        
        Returns:
            Path to generated file
        """
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_path = output_dir / "transcript.docx"
        
        doc = Document()
        
        # Title
        title = doc.add_heading("Meeting Transcript", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Metadata
        meta_para = doc.add_paragraph()
        meta_para.add_run(f"Duration: ").bold = True
        meta_para.add_run(format_timestamp(result.get("duration", 0)))
        meta_para.add_run("  |  ")
        meta_para.add_run(f"Generated: ").bold = True
        meta_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M"))
        
        # Speakers
        speakers = result.get("speakers", [])
        if speakers:
            doc.add_heading("Participants", level=1)
            for i, speaker in enumerate(speakers):
                para = doc.add_paragraph()
                colour = self.SPEAKER_COLOURS[i % len(self.SPEAKER_COLOURS)]
                
                name_run = para.add_run(f"{speaker.get('name', speaker.get('id', 'Unknown'))}")
                name_run.bold = True
                name_run.font.color.rgb = RGBColor(*colour)
                
                time_str = format_timestamp(speaker.get("speaking_time", 0))
                para.add_run(f" ({time_str} speaking time)")
        
        # Transcript
        doc.add_heading("Transcript", level=1)
        
        # Build speaker colour map
        speaker_colours = {}
        for i, speaker in enumerate(speakers):
            speaker_id = speaker.get("id", f"SPEAKER_{i}")
            speaker_name = speaker.get("name", speaker_id)
            speaker_colours[speaker_id] = self.SPEAKER_COLOURS[i % len(self.SPEAKER_COLOURS)]
            speaker_colours[speaker_name] = speaker_colours[speaker_id]
        
        # Track current speaker for grouping
        current_speaker = None
        current_para = None
        
        for seg in result.get("segments", []):
            speaker = seg.get("speaker_name", seg.get("speaker", "UNKNOWN"))
            text = seg.get("text", "").strip()
            
            if not text:
                continue
            
            # New paragraph for new speaker
            if speaker != current_speaker:
                current_speaker = speaker
                current_para = doc.add_paragraph()
                
                # Speaker name
                colour = speaker_colours.get(speaker, (0, 0, 0))
                name_run = current_para.add_run(f"{speaker}")
                name_run.bold = True
                name_run.font.color.rgb = RGBColor(*colour)
                
                # Timestamp
                if include_timestamps:
                    timestamp = format_timestamp(seg.get("start", 0))
                    ts_run = current_para.add_run(f" [{timestamp}]")
                    ts_run.font.size = Pt(9)
                    ts_run.font.color.rgb = RGBColor(128, 128, 128)
                
                current_para.add_run("\n")
            
            # Add text
            current_para.add_run(text + " ")
        
        # Save
        doc.save(output_path)
        
        logger.info(f"Generated DOCX: {output_path}")
        return output_path
    
    def generate_srt(self, result: dict, output_dir: Path) -> Path:
        """
        Generate SRT subtitle file.
        
        Args:
            result: Transcription result
            output_dir: Output directory
        
        Returns:
            Path to generated file
        """
        output_dir = Path(output_dir)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        output_path = output_dir / "transcript.srt"
        
        def format_srt_time(seconds: float) -> str:
            """Format time for SRT (HH:MM:SS,mmm)"""
            hours = int(seconds // 3600)
            minutes = int((seconds % 3600) // 60)
            secs = int(seconds % 60)
            millis = int((seconds % 1) * 1000)
            return f"{hours:02d}:{minutes:02d}:{secs:02d},{millis:03d}"
        
        lines = []
        for i, seg in enumerate(result.get("segments", []), 1):
            start = format_srt_time(seg.get("start", 0))
            end = format_srt_time(seg.get("end", 0))
            speaker = seg.get("speaker_name", seg.get("speaker", ""))
            text = seg.get("text", "").strip()
            
            if not text:
                continue
            
            lines.append(str(i))
            lines.append(f"{start} --> {end}")
            
            if speaker:
                lines.append(f"[{speaker}] {text}")
            else:
                lines.append(text)
            
            lines.append("")  # Blank line between entries
        
        with open(output_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        
        logger.info(f"Generated SRT: {output_path}")
        return output_path
